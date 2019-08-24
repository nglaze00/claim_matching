"""
Created by Nicholas Glaze

Module for reordering a local uspto claims document according to the US Patent & Trademark Office's
master copy. Uses separate documents for USPTO claims and local uspto claims.
"""



import docx, time, os
from collections import defaultdict


def slice(claim, prefix):
    """
    Removes all characters before a prefix in a claim string.
    :param claim: string
    :param prefix: string
    :return: claim w/ prefix removed
    """
    return claim.split(prefix, 1)[1]

def compare_strings(str1, str2):
    """
    Returns the difference in word counts between str1 and str2, based on str1.
    :param str1: string
    :param str2: string
    :return: int
    """
    diff = 0
    words1 = str1.split(" ")
    words2 = str2.split(" ")
    counts1, counts2 = defaultdict(int), defaultdict(int)

    #Compute word counts
    for word in words1:
        counts1[word] += 1
    for word in words2:
        counts2[word] += 1

    #Compute total difference
    diff = 0
    for word in counts1.keys():
        diff += abs(counts1[word] - counts2[word])

    return diff

def compute_diffs(key, other, slice, fast):
    """
    Computes the difference matrix for two lists of strings, according to compare_strings.
    :param key: list of strings
    :param other: list of strings
    :param slice: function to format strings before comparing
    :param fast: boolean; True: only computes scores for unmatched elements of other
    :return: diffs: matrix of ints where diffs[i, j] = compare_strings score between key[i] and other[j]; last index
                    in each row is index of min difference
    """
    reorder = [""] * len(other)
    matched = [False] * len(other)
    diffs = []  #Score matrix: row = scores for key

    #For each item in key, determine which item in other most closely matches it
    for index, item in enumerate(key):
        diffs_row = []
        closest = (None, float("inf"))
        for index, test in enumerate(other):
            if not matched[index] or not fast:
                diff = compare_strings(slice(item, "."), slice(test, "."))
                diffs_row.append(diff)
                if diff < closest[1]:
                    closest = (index, diff)
            else:
                diffs_row.append(float("inf"))
        matched[closest[0]] = True
        diffs_row.append(closest[0])
        diffs.append(diffs_row)
    return diffs

def reorder_lists(key, other, diffs):
    """
    Reorders the items of other to match those of key so that string differences are minimized
    :param key: list of strs
    :param other: list of strs
    :param diffs: difference matrix computed by compute_diffs
    :return: key and other, with other reordered to match key as closely as possible
    """
    reordered = [""] * len(other)
    for i in range(len(key)):
        if diffs[i][diffs[i][-1]] > 4:
            print("Claim " + str(i) + " may not match; check it!")
        reordered[i] = other[diffs[i][-1]]
    return key, reordered

def uspto_to_list(file):
    """
    Extracts and returns the claims present in a Google Patents file
    :param file: filename of uspto document
    :return: list of strings, each representing a claim
    """
    doc = docx.Document(file)
    paras = []
    after_header = False

    after_header, before_footer = False, True
    i = 0
    while i < len(doc.paragraphs):
        para = doc.paragraphs[i].text

        # Ignore document header and footer
        if after_header and para == "":
            before_footer = False
        if after_header and before_footer:
            paras.append(para)
            i += 1
        if para[:4] == "What":  # header: "What is claimed is:"
            after_header = True
            i += 1
        i += 1
    return paras

def local_to_list(file):
    """
    Extracts and returns the claims present in a local file
    :param file: filename of local document
    :return: list of strings, each representing a claim
    """
    doc = docx.Document(file)
    paras = []
    after_header = False

    after_header, before_footer = False, True
    i = 0
    para = ""
    while i < len(doc.paragraphs):

        item = doc.paragraphs[i].text
        if "(Canceled)" not in item:        # Ignore canceled claims
            item = item.replace("\n", "")   # Reformat without tabs / newlines
            item = item.replace("\t", " ")
            if len(item) > 1 and "." in item[1:5]:
                try:
                    paras.append(para)
                    para = ""

                except Exception as e:
                    para = ""
            para += " " + item

        i += 1
    paras.append(para)
    return paras[1:]

def reorder_local_claims(uspto_path, local_path):
    """
    Reorders the claims in local_file, using uspto_file as the key.
    :param uspto_path: string - filepath
    :param local_path: string - filepath
    :return: none; produces local_file_reordered.docx
    """
    other, key = local_to_list(local_path), uspto_to_list(uspto_path)
    assert len(key) > 0
    diffs = compute_diffs(key, other, slice, False)


    _, reordered = reorder_lists(key, other, diffs)
    doc = docx.Document(local_path)
    doc._body.clear_content()
    for i, claim in enumerate(reordered):
        doc.add_paragraph(claim)
        doc.add_paragraph("")

    filename = local_path.split("\\")[-1].split(".")[0] + "_reordered.docx"
    print("File " + filename + " created")

    doc.save("test\\local_reordered\\" + filename)


if __name__ == "__main__":
    try:
        try:
            os.mkdir("test\\local_reordered")
        except:
            pass
        uspto_dir = "test\\uspto"
        local_dir = "test\\local"
        uspto_files = os.listdir(uspto_dir)
        local_files = os.listdir(local_dir)

        assert len(uspto_files) == len(local_files), "# of USPTO and local files doesn't match"
        # Corresponding files must be at the same alphabetically sorted index
        # Recommend (uspto1.docx, local1.docx), (uspto2.docx, local2.docx), etc
        for uspto_file, local_file in zip(uspto_files, local_files):
            print("Matching " + local_file + " to " + uspto_file)z
            try:
                reorder_local_claims(uspto_dir + "\\" + uspto_file, local_dir + "\\" +  local_file)
            except Exception as e:
                print("Claim reordering failed:", e)


        print("Claims reordered successfully. This window will stay open until you close it.")
        time.sleep(10000)
    except Exception as e:
        print("Claim reordering failed:", e)
        time.sleep(10000)